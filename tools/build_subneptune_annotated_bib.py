from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("outputs")
OUT_DIR.mkdir(exist_ok=True)

TITLE = "Sub-Neptune Likely Connections: Proposal-Focused Annotated Bibliography"
SUBTITLE = (
    "Daniel Apai 2025 Sagan Workshop table, interpreted for reflected-light "
    "planet typing near the radius valley"
)


ORIENTATION = [
    (
        "Proposal lens",
        "The proposal asks whether radius-valley sub-Neptunes can masquerade as "
        "terrestrial worlds in low-resolution reflected-light spectra for HWO. "
        "The useful literature is therefore not only about what sub-Neptunes are, "
        "but about which physical processes move albedo, color, molecular bands, "
        "phase behavior, and radius/composition priors.",
    ),
    (
        "How to use the Apai table",
        "Treat each process branch as a modeling or interpretation knob: "
        "photochemistry changes absorbers and hazes; atmospheric loss and "
        "formation set envelope mass and radius; clouds and dynamics set "
        "continuum shape and phase-dependent albedo; outgassing, oceans, and "
        "magma oceans control secondary or hybrid atmospheric composition.",
    ),
    (
        "PICASO connection",
        "The most proposal-relevant papers help choose PICASO grid axes and "
        "interpret degeneracies: host spectrum and irradiation, metallicity, "
        "H2-rich versus secondary atmospheres, cloud-top pressure, cloud optical "
        "depth, phase angle, and wavelength regions that separate sub-Neptune "
        "spectra from terrestrial analogs.",
    ),
]


PROCESSES = [
    {
        "name": "Photochemistry",
        "property": "UV irradiation, effective temperature",
        "correlates": "Albedo, gas-phase abundances, thermal emission",
        "proposal_lens": (
            "Use this branch to decide which chemistry assumptions can change "
            "reflected-light colors, absorber depths, and haze/cloud opacity in "
            "PICASO-like spectra."
        ),
        "papers": [
            {
                "short": "Hu 2021",
                "citation": (
                    "Hu, R. (2021). Photochemistry and Spectral Characterization "
                    "of Temperate and Gas-rich Exoplanets. The Astrophysical "
                    "Journal, 921, 27."
                ),
                "link": "DOI: 10.3847/1538-4357/ac1789; arXiv:2108.04419",
                "role": (
                    "Sets the chemistry background for how UV fields and "
                    "temperature reshape gas-rich exoplanet spectra."
                ),
                "question": (
                    "How do disequilibrium reactions and photolysis alter the "
                    "observable atmospheres of temperate, gas-rich planets?"
                ),
                "sample": (
                    "Model study of H2-rich and gas-rich atmospheres with varied "
                    "stellar irradiation, temperature, and chemical pathways."
                ),
                "result": (
                    "Photochemistry can strongly alter CH4, NH3, H2O, sulfur, "
                    "and haze-related opacity, changing the apparent color and "
                    "molecular-band pattern even before clouds are considered."
                ),
                "proposal": (
                    "For your HWO/PICASO grid, this is the warning that the same "
                    "bulk sub-Neptune can occupy different reflected-light "
                    "locations depending on stellar UV and photochemical state. "
                    "It supports treating chemistry as a source of false "
                    "Earth-like or false non-Earth-like spectra."
                ),
                "caveat": (
                    "This is a model framework, not a direct radius-valley "
                    "population sample. Do not use it alone to infer occurrence "
                    "rates or radius-composition boundaries."
                ),
                "use": (
                    "Use for the chemistry rationale behind metallicity, UV, and "
                    "absorber assumptions in the proposal grid."
                ),
            },
            {
                "short": "Kempton et al. 2023",
                "citation": (
                    "Kempton, E. M.-R., et al. (2023). A reflective, metal-rich "
                    "atmosphere for GJ 1214b from its JWST phase curve. Nature, "
                    "620, 67-71."
                ),
                "link": "DOI: 10.1038/s41586-023-06159-5; arXiv:2305.06240",
                "role": (
                    "Empirical anchor for high-metallicity, reflective, cloudy "
                    "sub-Neptune atmospheres."
                ),
                "question": (
                    "What atmospheric structure explains the phase behavior and "
                    "brightness of the benchmark warm sub-Neptune GJ 1214b?"
                ),
                "sample": (
                    "JWST phase-curve observation of GJ 1214b, interpreted with "
                    "atmospheric models that test metallicity, clouds/hazes, and "
                    "heat redistribution."
                ),
                "result": (
                    "The planet is consistent with a reflective atmosphere with "
                    "high metallicity and strong aerosol/cloud influence rather "
                    "than a simple clear solar-composition H2 envelope."
                ),
                "proposal": (
                    "This is one of the most important proposal papers because it "
                    "shows that a sub-Neptune can be bright and reflective. In "
                    "reflected light, that means radius and surface type cannot "
                    "be inferred from flux ratio alone; clouds and metallicity "
                    "can help an un-Earth look deceptively terrestrial."
                ),
                "caveat": (
                    "GJ 1214b is a transiting warm sub-Neptune, not an HWO "
                    "direct-imaging target. Use it as a physical analog, not a "
                    "direct observing-mode analog."
                ),
                "use": (
                    "Use for the case that high metallicity and aerosols belong "
                    "in the PICASO grid and in the planet-typing discussion."
                ),
            },
            {
                "short": "Reed et al. 2024",
                "citation": (
                    "Reed, N. W., Shearer, R. L., McGlynn, S. E., et al. (2024). "
                    "Abiotic Production of Dimethyl Sulfide, Carbonyl Sulfide, "
                    "and Other Organosulfur Gases via Photochemistry: "
                    "Implications for Biosignatures and Metabolic Potential. "
                    "The Astrophysical Journal Letters, 973, L38."
                ),
                "link": "DOI: 10.3847/2041-8213/ad74da",
                "role": (
                    "Shows that photochemistry can generate species often "
                    "treated as biologically suggestive or atmospherically "
                    "diagnostic."
                ),
                "question": (
                    "Can organosulfur gases arise abiotically through "
                    "photochemical pathways?"
                ),
                "sample": (
                    "Photochemical and geochemical context for sulfur-bearing "
                    "molecules, focused on abiotic production pathways and their "
                    "spectral/interpretive consequences."
                ),
                "result": (
                    "Potential biosignature-like sulfur gases can be produced "
                    "without biology under some photochemical conditions."
                ),
                "proposal": (
                    "For the proposal, this is less about a main PICASO opacity "
                    "axis and more about interpretation discipline: spectral "
                    "features and colors can have abiotic chemical origins, so "
                    "planet typing should not over-read a single molecular cue."
                ),
                "caveat": (
                    "The paper is not a sub-Neptune occurrence or direct-imaging "
                    "grid paper; keep it as a photochemical interpretive caveat."
                ),
                "use": (
                    "Use for the literature-review sentence that chemistry can "
                    "create misleading atmospheric diagnostics."
                ),
            },
        ],
    },
    {
        "name": "Atmospheric loss",
        "property": "Planet size, density",
        "correlates": "Age, irradiation, core mass, composition",
        "proposal_lens": (
            "Use this branch to connect the radius valley to envelope mass, "
            "age, irradiation, and why radius cannot be assumed for directly "
            "imaged targets."
        ),
        "papers": [
            {
                "short": "Owen and Wu 2017",
                "citation": (
                    "Owen, J. E., and Wu, Y. (2017). The Evaporation Valley in "
                    "the Kepler Planets. The Astrophysical Journal, 847, 29."
                ),
                "link": "DOI: 10.3847/1538-4357/aa890a; arXiv:1705.10810",
                "role": (
                    "Photoevaporation model for the radius valley and envelope "
                    "loss."
                ),
                "question": (
                    "Can stellar high-energy irradiation carve the observed "
                    "radius gap by stripping H/He envelopes?"
                ),
                "sample": (
                    "Theoretical population/evolution modeling tied to the Kepler "
                    "small-planet radius distribution."
                ),
                "result": (
                    "Atmospheric escape can produce a valley between bare rocky "
                    "cores and planets retaining volatile envelopes, with the "
                    "gap depending on irradiation and core/envelope properties."
                ),
                "proposal": (
                    "This gives your project the physical reason the radius "
                    "valley matters: it is a composition boundary shaped by "
                    "time-dependent mass loss. For HWO, where radius is not "
                    "measured directly, spectra must carry some of that typing "
                    "burden."
                ),
                "caveat": (
                    "It is not a reflected-light paper. Use it to justify the "
                    "population boundary, not to predict albedo or spectral bands."
                ),
                "use": (
                    "Use for the atmospheric-loss explanation of why near-valley "
                    "sub-Neptunes are the dangerous impostor class."
                ),
            },
            {
                "short": "Fulton et al. 2017",
                "citation": (
                    "Fulton, B. J., et al. (2017). The California-Kepler Survey. "
                    "III. A Gap in the Radius Distribution of Small Planets. "
                    "The Astronomical Journal, 154, 109."
                ),
                "link": "DOI: 10.3847/1538-3881/aa80eb; arXiv:1611.09137",
                "role": (
                    "Empirical discovery/measurement of the radius gap in a "
                    "high-precision Kepler stellar/planet sample."
                ),
                "question": (
                    "What does the occurrence distribution of small planets look "
                    "like once stellar radii are improved?"
                ),
                "sample": (
                    "California-Kepler Survey planets with improved stellar "
                    "parameters, enabling a sharper small-planet radius "
                    "distribution."
                ),
                "result": (
                    "The survey finds a deficit of planets near roughly "
                    "1.5-2.0 R_Earth, separating smaller super-Earths from "
                    "larger sub-Neptunes."
                ),
                "proposal": (
                    "This is the empirical backbone of your introduction. It "
                    "defines why a planet just above the valley is not simply a "
                    "large Earth, and why direct-imaging missions need a "
                    "spectroscopic typing strategy."
                ),
                "caveat": (
                    "The result comes from transiting planets. HWO targets will "
                    "not usually have transit radii, so the paper motivates the "
                    "problem rather than solving the direct-imaging degeneracy."
                ),
                "use": (
                    "Use for the first paragraph establishing the radius valley "
                    "and the risk of radius-free planet typing."
                ),
            },
            {
                "short": "Gupta and Schlichting 2019",
                "citation": (
                    "Gupta, A., and Schlichting, H. E. (2019). Sculpting the "
                    "valley in the radius distribution of small exoplanets as a "
                    "by-product of planet formation: the core-powered mass-loss "
                    "mechanism. Monthly Notices of the Royal Astronomical "
                    "Society, 487, 24-33."
                ),
                "link": "DOI: 10.1093/mnras/stz1230; arXiv:1811.03202",
                "role": (
                    "Alternative atmospheric-loss mechanism driven by planetary "
                    "cooling rather than only stellar XUV irradiation."
                ),
                "question": (
                    "Can planets lose envelopes because core cooling powers "
                    "escape, even without requiring extreme XUV-driven removal?"
                ),
                "sample": (
                    "Theoretical population modeling of small planets, envelope "
                    "cooling, and mass-loss evolution."
                ),
                "result": (
                    "Core-powered mass loss can reproduce the radius valley and "
                    "links final radius to core mass, equilibrium temperature, "
                    "and primordial envelope properties."
                ),
                "proposal": (
                    "This keeps your proposal from sounding like the radius "
                    "valley has only one cause. For reflected-light typing, the "
                    "important point is that several evolutionary histories can "
                    "produce similar radii but different atmospheric states."
                ),
                "caveat": (
                    "It is a population/evolution model, not a spectra paper. "
                    "Use it to frame priors, not to set spectral opacities."
                ),
                "use": (
                    "Use for a balanced atmospheric-loss literature review and "
                    "for explaining age/temperature sensitivity."
                ),
            },
            {
                "short": "Fernandes et al. 2025",
                "citation": (
                    "Fernandes, R. B., et al. (2025). Signatures of Atmospheric "
                    "Mass Loss and Planet Migration in the Time Evolution of "
                    "Short-period Transiting Exoplanets. The Astronomical "
                    "Journal, 169, 208."
                ),
                "link": "DOI: 10.3847/1538-3881/adb97e; arXiv:2503.10856",
                "role": (
                    "Young versus old planet populations as evidence for "
                    "time-dependent loss and migration."
                ),
                "question": (
                    "How does the occurrence of short-period planets evolve with "
                    "stellar age, and what does that imply for atmospheric mass "
                    "loss and migration?"
                ),
                "sample": (
                    "TESS-based occurrence study of short-period 1.8-10 R_Earth "
                    "planets around nearby young FGK cluster stars, compared to "
                    "older field populations."
                ),
                "result": (
                    "Young populations show a different occurrence pattern, "
                    "supporting the view that mass loss and orbital evolution "
                    "change the visible sub-Neptune population over time."
                ),
                "proposal": (
                    "This is useful for connecting your static spectral grid to "
                    "evolution. A planet's reflected-light spectrum is observed "
                    "at one time, but its radius, envelope fraction, and "
                    "composition are products of age, loss, and migration."
                ),
                "caveat": (
                    "The sample is short-period transiting planets, not HZ "
                    "direct-imaging targets. Use it for evolutionary context, "
                    "not as the HWO target distribution."
                ),
                "use": (
                    "Use for the proposal argument that age and irradiation "
                    "should appear in the interpretation of near-valley planets."
                ),
            },
        ],
    },
    {
        "name": "Condensate clouds",
        "property": "Effective temperature, circulation",
        "correlates": "Infrared color, albedo",
        "proposal_lens": (
            "Use this branch to explain how clouds and aerosols can flatten "
            "bands, raise albedo, and erase otherwise diagnostic spectral shape."
        ),
        "papers": [
            {
                "short": "Hu 2019",
                "citation": (
                    "Hu, R. (2019). Information in the Reflected-light Spectra "
                    "of Widely Separated Giant Exoplanets. The Astrophysical "
                    "Journal, 887, 166."
                ),
                "link": "DOI: 10.3847/1538-4357/ab58c7",
                "role": (
                    "Reflected-light information-content paper for cloud, gas, "
                    "albedo, and color retrieval logic."
                ),
                "question": (
                    "What physical information can low- to moderate-resolution "
                    "reflected-light spectra actually recover?"
                ),
                "sample": (
                    "Forward-model and information-content analysis of directly "
                    "imaged, widely separated giant exoplanets."
                ),
                "result": (
                    "Reflected-light spectra encode clouds, absorbers, and "
                    "albedo, but the information is limited and often "
                    "degenerate across atmospheric properties."
                ),
                "proposal": (
                    "Even though the target class is giants, the observing mode "
                    "is close to your HWO problem. This paper helps justify why "
                    "your project should compare full spectral shapes and not "
                    "single flux ratios."
                ),
                "caveat": (
                    "It is not a sub-Neptune paper. Use it for reflected-light "
                    "methodology and degeneracy language, not sub-Neptune "
                    "interior priors."
                ),
                "use": (
                    "Use for the reflected-light retrieval/degeneracy framework "
                    "and as a bridge from PICASO spectra to HWO observables."
                ),
            }
        ],
    },
    {
        "name": "Atmosphere-ocean coupling",
        "property": "H2O abundance, effective temperature",
        "correlates": "Planet radius, density, irradiation",
        "proposal_lens": (
            "Use this branch to keep the terrestrial-comparison set honest: "
            "oceans, shallow surfaces, and volatile exchange can make secondary "
            "and hybrid atmospheres behave differently from simple gas envelopes."
        ),
        "papers": [
            {
                "short": "Hu and Delgado Diaz 2019",
                "citation": (
                    "Hu, R., and Delgado Diaz, B. (2019). Stability of Nitrogen "
                    "in Planetary Atmospheres in Contact with Liquid Water. The "
                    "Astrophysical Journal, 886, 126."
                ),
                "link": "DOI: 10.3847/1538-4357/ab4cea",
                "role": (
                    "Atmosphere-ocean chemistry constraint on the persistence "
                    "of N2 in water-contact atmospheres."
                ),
                "question": (
                    "Can an atmosphere in contact with liquid water retain "
                    "nitrogen over long timescales?"
                ),
                "sample": (
                    "Thermodynamic and geochemical modeling of planetary "
                    "atmospheres interacting with liquid-water reservoirs."
                ),
                "result": (
                    "Ocean-atmosphere contact can change volatile inventories "
                    "and atmospheric stability, including the fate of nitrogen."
                ),
                "proposal": (
                    "Your comparison set includes Earth and abiotic terrestrial "
                    "analogs. This paper is a reminder that secondary "
                    "atmospheres are regulated by reservoirs, not just by "
                    "top-of-atmosphere irradiation."
                ),
                "caveat": (
                    "It is not a sub-Neptune spectral grid paper. Its value is "
                    "composition logic for terrestrial analogs and shallow-water "
                    "cases."
                ),
                "use": (
                    "Use for the caveat that Earth-like reflected-light spectra "
                    "depend on atmosphere-surface reservoirs."
                ),
            },
            {
                "short": "Tsai et al. 2021",
                "citation": (
                    "Tsai, S.-M., Innes, H., Lichtenberg, T., et al. (2021). "
                    "Inferring Shallow Surfaces on Sub-Neptune Exoplanets with "
                    "JWST. The Astrophysical Journal Letters, 922, L27."
                ),
                "link": "DOI: 10.3847/2041-8213/ac399a; arXiv:2109.12268",
                "role": (
                    "Connects sub-Neptune atmospheric observations to the "
                    "possibility of shallow surfaces or bounded atmospheres."
                ),
                "question": (
                    "Can observations distinguish a deep gas envelope from a "
                    "sub-Neptune atmosphere limited by a shallow surface?"
                ),
                "sample": (
                    "Modeling study of sub-Neptune atmospheric scenarios and "
                    "JWST-observable signatures."
                ),
                "result": (
                    "Some atmospheric/spectral patterns can reveal whether a "
                    "sub-Neptune has a shallow boundary rather than an extended "
                    "deep envelope."
                ),
                "proposal": (
                    "This is directly relevant to the masquerade problem: a "
                    "planet can be larger than Earth but have atmospheric or "
                    "surface constraints that complicate a simple gas-world "
                    "classification."
                ),
                "caveat": (
                    "The observing mode is JWST transmission/emission context, "
                    "not HWO reflected light. Translate the principle, not the "
                    "instrument-specific detectability."
                ),
                "use": (
                    "Use for the argument that sub-Neptunes may span gas-rich, "
                    "surface-limited, and hybrid atmospheric regimes."
                ),
            },
        ],
    },
    {
        "name": "Migration/accretion",
        "property": "Semi-major axis, irradiation",
        "correlates": "Envelope mass, density, age, orbital architecture",
        "proposal_lens": (
            "Use this branch to convert formation history into reasonable "
            "composition priors for the spectral grid: water inventory, H/He "
            "envelope mass, and where planets formed relative to the ice line."
        ),
        "papers": [
            {
                "short": "Johansen and Lambrechts 2017",
                "citation": (
                    "Johansen, A., and Lambrechts, M. (2017). Forming Planets "
                    "via Pebble Accretion. Annual Review of Earth and Planetary "
                    "Sciences, 45, 359-387."
                ),
                "link": "DOI: 10.1146/annurev-earth-063016-020226",
                "role": (
                    "Review of pebble accretion as a fast route to planetary "
                    "cores and volatile-rich planets."
                ),
                "question": (
                    "How can small solids drifting through disks build cores "
                    "efficiently enough to seed planets?"
                ),
                "sample": (
                    "Review synthesis of pebble-accretion theory, disk dynamics, "
                    "growth barriers, and planet formation regimes."
                ),
                "result": (
                    "Pebble accretion can rapidly form cores whose final "
                    "composition and envelope potential depend on disk location, "
                    "solid flux, and migration history."
                ),
                "proposal": (
                    "This is a field-entry paper: it explains why sub-Neptunes "
                    "are expected to have diverse cores and volatile inventories, "
                    "which becomes the prior behind the PICASO metallicity and "
                    "water-rich cases."
                ),
                "caveat": (
                    "It is a broad review, not a spectra paper or one specific "
                    "population model."
                ),
                "use": (
                    "Use for the formation background that connects orbital "
                    "architecture to composition diversity."
                ),
            },
            {
                "short": "Bitsch et al. 2021",
                "citation": (
                    "Bitsch, B., et al. (2021). Dry or water world? How the "
                    "water contents of inner sub-Neptunes constrain giant planet "
                    "formation and the location of the water ice line. Astronomy "
                    "and Astrophysics, 649, L5."
                ),
                "link": "DOI: 10.1051/0004-6361/202140793",
                "role": (
                    "Formation-composition link for water-rich versus dry inner "
                    "sub-Neptunes."
                ),
                "question": (
                    "What can the water content of inner sub-Neptunes tell us "
                    "about where they formed and how disks evolved?"
                ),
                "sample": (
                    "Planet-formation modeling tied to ice-line location, giant "
                    "planet formation, and inward delivery of volatile-rich "
                    "material."
                ),
                "result": (
                    "Inner sub-Neptune water content can encode formation beyond "
                    "or inside the water ice line and the timing of giant-planet "
                    "formation."
                ),
                "proposal": (
                    "This paper helps you justify water-rich and high-metallicity "
                    "cases in the reflected-light grid. The same radius can hide "
                    "different interior water fractions, which can feed back into "
                    "atmospheric mean molecular weight and spectral appearance."
                ),
                "caveat": (
                    "The conclusions depend on disk and migration assumptions. "
                    "Do not treat water content as uniquely recoverable from one "
                    "spectrum."
                ),
                "use": (
                    "Use for the formation rationale behind including water-rich "
                    "sub-Neptune analogs."
                ),
            },
            {
                "short": "Ida and Lin 2010",
                "citation": (
                    "Ida, S., and Lin, D. N. C. (2010). Toward a Deterministic "
                    "Model of Planetary Formation. VI. Dynamical Interaction and "
                    "Coagulation of Multiple Rocky Embryos and Super-Earth "
                    "Systems around Solar-Type Stars. The Astrophysical Journal, "
                    "719, 810-830."
                ),
                "link": "DOI: 10.1088/0004-637X/719/1/810; arXiv:1006.2584",
                "role": (
                    "Early deterministic formation model for rocky embryos and "
                    "super-Earth architectures."
                ),
                "question": (
                    "How do multiple embryos grow, interact, and assemble "
                    "compact super-Earth systems around Sun-like stars?"
                ),
                "sample": (
                    "Semi-analytic and numerical formation modeling of embryo "
                    "growth, interaction, and orbital architecture."
                ),
                "result": (
                    "Planet formation naturally produces a range of core masses, "
                    "semi-major axes, and architectures that later shape "
                    "envelope acquisition and loss."
                ),
                "proposal": (
                    "For your work, this is background for why semi-major axis "
                    "and architecture appear in Apai's correlate column. Those "
                    "formation variables become hidden priors behind the "
                    "reflected-light spectrum."
                ),
                "caveat": (
                    "It predates much of the modern radius-valley literature and "
                    "does not model HWO spectra."
                ),
                "use": (
                    "Use for historical formation context and the link between "
                    "architecture and planet type."
                ),
            },
            {
                "short": "Rogers et al. 2011",
                "citation": (
                    "Rogers, L. A., Bodenheimer, P., Lissauer, J. J., and Seager, "
                    "S. (2011). Formation and Structure of Low-Density "
                    "Exo-Neptunes. The Astrophysical Journal, 738, 59."
                ),
                "link": "DOI: 10.1088/0004-637X/738/1/59; arXiv:1106.2807",
                "role": (
                    "Interior/envelope structure paper for low-density "
                    "sub-Neptune and exo-Neptune planets."
                ),
                "question": (
                    "What combinations of core, water, and H/He envelope can "
                    "produce low-density exo-Neptune radii?"
                ),
                "sample": (
                    "Planet interior and formation modeling for low-density "
                    "exo-Neptunes with varied composition and envelope fractions."
                ),
                "result": (
                    "Small H/He envelopes or volatile layers can substantially "
                    "inflate radius, producing large radius-composition "
                    "degeneracy."
                ),
                "proposal": (
                    "This is central to the masquerade idea: radius alone is not "
                    "composition, and composition alone is not spectrum. It "
                    "supports treating sub-Neptunes as a family of envelope and "
                    "volatile states rather than one template."
                ),
                "caveat": (
                    "The paper focuses on structure and formation, not reflected "
                    "light."
                ),
                "use": (
                    "Use for explaining why a small volatile envelope can make a "
                    "planet observationally non-terrestrial."
                ),
            },
        ],
    },
    {
        "name": "Outgassing from core/mantle",
        "property": "Atmospheric composition",
        "correlates": "Planet radius, density, irradiation, stellar abundances",
        "proposal_lens": (
            "Use this branch to bridge rocky/interior chemistry and secondary "
            "or hybrid atmospheric composition in the comparison set."
        ),
        "papers": [
            {
                "short": "Grewal et al. 2021",
                "citation": (
                    "Grewal, D. S., et al. (2021). The effect of carbon "
                    "concentration on its core-mantle partitioning behavior in "
                    "inner Solar System rocky bodies. Earth and Planetary "
                    "Science Letters, 571, 117090."
                ),
                "link": "DOI: 10.1016/j.epsl.2021.117090",
                "role": (
                    "Geochemical constraint on carbon partitioning, interior "
                    "inventory, and later outgassing potential."
                ),
                "question": (
                    "How does carbon concentration affect whether carbon is "
                    "stored in cores or mantles during rocky-body differentiation?"
                ),
                "sample": (
                    "High-pressure geochemistry/partitioning study relevant to "
                    "rocky-body volatile budgets."
                ),
                "result": (
                    "Carbon storage depends on differentiation conditions, which "
                    "controls how much carbon remains available for mantle and "
                    "atmospheric reservoirs."
                ),
                "proposal": (
                    "This helps connect interior composition to possible CO2, "
                    "CO, or CH4-rich secondary atmospheres. It is useful when "
                    "arguing that a planet's spectrum reflects both atmosphere "
                    "and interior history."
                ),
                "caveat": (
                    "It is Solar System geochemistry, not a sub-Neptune spectrum "
                    "paper. Use it as a mechanism reference, not a direct model "
                    "input."
                ),
                "use": (
                    "Use for the outgassing/interior-chemistry context behind "
                    "secondary atmosphere analogs."
                ),
            },
            {
                "short": "Tian and Heng 2024",
                "citation": (
                    "Tian, M., and Heng, K. (2024). Atmospheric Chemistry of "
                    "Secondary and Hybrid Atmospheres of Super Earths and "
                    "Sub-Neptunes. The Astrophysical Journal, 963, 157."
                ),
                "link": "DOI: 10.3847/1538-4357/ad217c; arXiv:2301.10217",
                "role": (
                    "Direct chemistry framework for secondary and hybrid "
                    "super-Earth/sub-Neptune atmospheres."
                ),
                "question": (
                    "How do outgassed species behave chemically when mixed with "
                    "H2-rich or hybrid envelopes?"
                ),
                "sample": (
                    "Atmospheric chemistry models spanning secondary and hybrid "
                    "composition regimes for super-Earths and sub-Neptunes."
                ),
                "result": (
                    "Hybrid atmospheres can have chemical states distinct from "
                    "both pure rocky secondary atmospheres and clear H2-rich "
                    "mini-Neptunes."
                ),
                "proposal": (
                    "This is highly relevant to your grid because the dangerous "
                    "impostor may not be a simple H2 envelope. Hybrid chemistry "
                    "can change H2O, CO2, CH4, and continuum opacity in ways "
                    "that affect planet typing."
                ),
                "caveat": (
                    "Chemical outputs still need translation through "
                    "reflected-light radiative transfer and cloud assumptions."
                ),
                "use": (
                    "Use for defining secondary/hybrid atmospheric cases in the "
                    "PICASO study."
                ),
            },
            {
                "short": "Kite et al. 2019",
                "citation": (
                    "Kite, E. S., Fegley, B., Schaefer, L., and Ford, E. B. "
                    "(2019). Superabundance of Exoplanet Sub-Neptunes Explained "
                    "by Fugacity Crisis. The Astrophysical Journal Letters, 887, "
                    "L33."
                ),
                "link": "DOI: 10.3847/2041-8213/ab59d9",
                "role": (
                    "Links interior redox/outgassing physics to the abundance of "
                    "sub-Neptune-like planets."
                ),
                "question": (
                    "Can magma/interior chemistry and volatile fugacity help "
                    "explain why sub-Neptunes are so common?"
                ),
                "sample": (
                    "Interior-atmosphere theory connecting volatile storage, "
                    "redox state, and planet radius outcomes."
                ),
                "result": (
                    "Interior volatile behavior can favor atmospheres/envelopes "
                    "that help create abundant sub-Neptune radii."
                ),
                "proposal": (
                    "This paper is useful because it pushes the proposal beyond "
                    "surface-level atmosphere models. It says some spectral "
                    "states may be consequences of interior-atmosphere coupling, "
                    "not just present-day irradiation."
                ),
                "caveat": (
                    "It is a theory paper and does not predict HWO spectra "
                    "directly."
                ),
                "use": (
                    "Use for the bridge between outgassing, interior evolution, "
                    "and why sub-Neptunes are common enough to matter for HWO."
                ),
            },
        ],
    },
    {
        "name": "Coupling to magma oceans",
        "property": "Atmospheric composition, transit radius",
        "correlates": "Effective temperature, density, irradiation, age",
        "proposal_lens": (
            "Use this branch to handle volatile sequestration, runaway "
            "greenhouse, and demographic transitions that affect the boundary "
            "between rocky planets and volatile-rich impostors."
        ),
        "papers": [
            {
                "short": "Dorn and Lichtenberg 2021",
                "citation": (
                    "Dorn, C., and Lichtenberg, T. (2021). Hidden Water in Magma "
                    "Ocean Exoplanets. The Astrophysical Journal Letters, 922, "
                    "L4."
                ),
                "link": "DOI: 10.3847/2041-8213/ac33af; arXiv:2110.15069",
                "role": (
                    "Magma-ocean/interior storage paper for hidden water and "
                    "volatile partitioning."
                ),
                "question": (
                    "How much water can be hidden in magma oceans and interiors "
                    "rather than appearing as an atmosphere or ocean?"
                ),
                "sample": (
                    "Interior-structure and volatile-partitioning modeling of "
                    "magma-ocean exoplanets."
                ),
                "result": (
                    "Large volatile inventories can be hidden in interiors, "
                    "changing the relation between bulk composition, radius, and "
                    "observable atmosphere."
                ),
                "proposal": (
                    "This matters because your project compares spectra, not "
                    "interiors. A planet's reflected-light appearance may hide "
                    "a very different water inventory than simple radius or "
                    "density intuition suggests."
                ),
                "caveat": (
                    "It is not a reflected-light spectral study. Use it for "
                    "composition priors and hidden-reservoir caveats."
                ),
                "use": (
                    "Use for explaining why water-rich sub-Neptunes may not have "
                    "obvious water-rich spectra."
                ),
            },
            {
                "short": "Schlecker et al. 2024",
                "citation": (
                    "Schlecker, M., Apai, D., Lichtenberg, T., et al. (2024). "
                    "Bioverse: The Habitable Zone Inner Edge Discontinuity as an "
                    "Imprint of Runaway Greenhouse Climates on Exoplanet "
                    "Demographics. The Planetary Science Journal, 5, 3."
                ),
                "link": "DOI: 10.3847/PSJ/acf57f; arXiv:2309.04518",
                "role": (
                    "Population-level link between climate transitions and "
                    "exoplanet demographics."
                ),
                "question": (
                    "Can runaway greenhouse climates leave detectable "
                    "demographic signatures near the habitable-zone inner edge?"
                ),
                "sample": (
                    "Bioverse simulation framework connecting climate outcomes "
                    "to observable exoplanet population patterns."
                ),
                "result": (
                    "Climate transitions can imprint population discontinuities, "
                    "suggesting that demographics and spectra should be read "
                    "together."
                ),
                "proposal": (
                    "This helps your proposal connect HWO target typing to a "
                    "larger statistical program: one spectrum is ambiguous, but "
                    "spectral classifications across a population can reveal "
                    "where volatile-rich or runaway states dominate."
                ),
                "caveat": (
                    "The paper is demographic and simulation-based; it does not "
                    "supply PICASO spectra."
                ),
                "use": (
                    "Use for the future-work argument that your spectral grid can "
                    "feed population-level planet classification."
                ),
            },
        ],
    },
    {
        "name": "Atmospheric dynamics",
        "property": "Irradiation, rotation period",
        "correlates": "Pressure-dependent albedo, intensity, color evolution",
        "proposal_lens": (
            "Use this branch to connect phase angle, heat redistribution, "
            "cloud location, and rotation to the reflected-light observables."
        ),
        "papers": [
            {
                "short": "Innes and Pierrehumbert 2022",
                "citation": (
                    "Innes, H., and Pierrehumbert, R. T. (2022). Atmospheric "
                    "Dynamics of Temperate Sub-Neptunes. I. Dry Dynamics. The "
                    "Astrophysical Journal, 927, 38."
                ),
                "link": "DOI: 10.3847/1538-4357/ac4887",
                "role": (
                    "Dry GCM baseline for temperate, low-mean-molecular-weight "
                    "sub-Neptune atmospheres."
                ),
                "question": (
                    "How do irradiation and rotation shape circulation in "
                    "temperate sub-Neptune atmospheres?"
                ),
                "sample": (
                    "Idealized ExoFMS GCM simulations of H2-rich temperate "
                    "sub-Neptunes across instellation and rotation regimes."
                ),
                "result": (
                    "Circulation regimes vary with rotation and irradiation, "
                    "affecting temperature patterns that can influence clouds, "
                    "phase curves, and emergent spectra."
                ),
                "proposal": (
                    "For reflected-light HWO work, this is the bridge from a "
                    "one-dimensional atmosphere to a disk-integrated planet. "
                    "Phase angle and atmospheric dynamics can change apparent "
                    "color and albedo, so spectra should not be treated as "
                    "static global averages without care."
                ),
                "caveat": (
                    "The simulations are dry and idealized, so cloud microphysics "
                    "and full chemistry still need separate treatment."
                ),
                "use": (
                    "Use for justifying phase angle and circulation sensitivity "
                    "as part of the PICASO/HWO interpretation."
                ),
            }
        ],
    },
]


def field_pairs(paper):
    return [
        ("Full citation and link", f"{paper['citation']} {paper['link']}."),
        ("Apai-table role", paper["role"]),
        ("Core question", paper["question"]),
        ("Sample/model setup", paper["sample"]),
        ("Key result", paper["result"]),
        ("Why it matters for your proposal", paper["proposal"]),
        ("Caveats / do not overclaim", paper["caveat"]),
        ("Use this paper for", paper["use"]),
    ]


def build_xmind_markdown():
    lines = [f"# {TITLE}", ""]
    lines.append(f"- {SUBTITLE}")
    lines.append("- Field orientation for the proposal")
    for label, text in ORIENTATION:
        lines.append(f"  - {label}: {text}")
    for process in PROCESSES:
        lines.append(f"- {process['name']}")
        lines.append(f"  - Property: {process['property']}")
        lines.append(f"  - Likely correlates: {process['correlates']}")
        lines.append(f"  - Proposal lens: {process['proposal_lens']}")
        for paper in process["papers"]:
            lines.append(f"  - {paper['short']}")
            for label, text in field_pairs(paper):
                lines.append(f"    - {label}: {text}")
    return "\n".join(lines) + "\n"


def build_readable_markdown():
    lines = [f"# {TITLE}", "", SUBTITLE, ""]
    lines.append("## Field orientation for the proposal")
    for label, text in ORIENTATION:
        lines.append(f"**{label}.** {text}")
        lines.append("")
    for process in PROCESSES:
        lines.append(f"## {process['name']}")
        lines.append(f"**Property.** {process['property']}")
        lines.append(f"**Likely correlates.** {process['correlates']}")
        lines.append(f"**Proposal lens.** {process['proposal_lens']}")
        lines.append("")
        for paper in process["papers"]:
            lines.append(f"### {paper['short']}")
            for label, text in field_pairs(paper):
                lines.append(f"**{label}.** {text}")
                lines.append("")
    return "\n".join(lines)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = "Table Text"
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "D9E2EC")


def set_table_geometry(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")


def configure_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Table Text" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Table Text", 1)
    else:
        style = doc.styles["Table Text"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    run = footer.add_run("Sub-Neptune Annotated Bibliography")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(f"{label}. ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    p.add_run(text)


def add_process_summary_table(doc, process):
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [1.5, 5.0])
    set_table_borders(table)
    rows = [
        ("Property", process["property"]),
        ("Likely correlates", process["correlates"]),
        ("Proposal lens", process["proposal_lens"]),
    ]
    for row, (label, text) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        set_cell_text(row.cells[0], label, bold=True, color="1F4D78")
        set_cell_text(row.cells[1], text)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    configure_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    srun = subtitle.add_run(SUBTITLE)
    srun.font.size = Pt(11)
    srun.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("Field orientation for the proposal", level=1)
    for label, text in ORIENTATION:
        add_label_paragraph(doc, label, text)

    doc.add_heading("Annotated bibliography by Apai process", level=1)
    for idx, process in enumerate(PROCESSES):
        if idx:
            doc.add_page_break()
        doc.add_heading(process["name"], level=2)
        add_process_summary_table(doc, process)
        for paper in process["papers"]:
            doc.add_heading(paper["short"], level=3)
            for label, text in field_pairs(paper):
                add_label_paragraph(doc, label, text)

    doc.add_heading("Source and collection notes", level=1)
    notes = [
        (
            "Source table",
            "Daniel Apai 2025 Sagan Workshop slide deck, Likely Connections "
            "table. The bibliography is organized by the process rows in that "
            "table rather than by publication year.",
        ),
        (
            "Zotero state",
            "The target collection contains 20 parent papers plus a corrected "
            "note. Fernandes et al. 2025 replaces the earlier 2023 fallback for "
            "the Fernandez+25 shorthand.",
        ),
        (
            "PDF policy",
            "Only legal/open PDFs were attached or referenced. Zotero full-text "
            "lookup did not find files for Johansen and Lambrechts 2017 or "
            "Grewal et al. 2021.",
        ),
    ]
    for label, text in notes:
        add_label_paragraph(doc, label, text)

    output = OUT_DIR / "Sub_Neptune_Apai_Annotated_Bibliography.docx"
    doc.save(output)
    return output


def main():
    xmind_md = OUT_DIR / "Sub_Neptune_Apai_Annotated_Bibliography_XMind.md"
    readable_md = OUT_DIR / "Sub_Neptune_Apai_Annotated_Bibliography.md"
    xmind_md.write_text(build_xmind_markdown(), encoding="utf-8")
    readable_md.write_text(build_readable_markdown(), encoding="utf-8")
    docx = build_docx()
    print(xmind_md)
    print(readable_md)
    print(docx)


if __name__ == "__main__":
    main()
